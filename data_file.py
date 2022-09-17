# Imported os library for creating the files, to update the file and to retrieve the data from it.
# docx allows us to access files saved in doc format whenever you want
# Note: Instead of docx install python-docx


from docx import Document
from website1 import weather_data
from datetime import datetime
import os

doc = Document()  # Create an instance of word document
doc.add_heading('Weather Forecasting⛈️🌤️🌦️🌈🌞', 0)  # Add a title to the document
date = datetime.now().date()
time = datetime.now().time().strftime("%H:%M:%S")
doc.add_heading(
    f'                                                                                       Today\'s Date: {date} \n        '
    f'                                                                                                Time: {time}', 1)


def data():
    x = weather_data()
    if x is not None:
        location = x['location']  # To extract the data that is stored in a variable from weather file
        time = x['time']
        information = x['information']
        weather = x['weather']
        humidity = x['humidity']
        wind_speed = x['wind_speed']
        precipitation = x['Precipitation']

        # To print the result on word file
        doc.add_paragraph(
            f'Location: {location}\nTime: {time}\nWeather: {information}\nTemperature: {weather}°C\nHumidity: {humidity}\n'
            f'Wind Speed: {wind_speed}\nPrecipitation: {precipitation}')
    else:
        pass


def file():
    file_name = input("Enter the file name for what you have searched till now: ")
    doc.save(f"{file_name}.docx")  # To save the Word document
    os.system(f"{file_name}.docx")  # To open the document in the output itself
    quit()
